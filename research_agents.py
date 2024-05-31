import os
from crewai import Agent, Task, Crew
from utils import get_openai_api_key
from docx import Document
import openai


from dotenv import load_dotenv
load_dotenv()
openai.api_key = os.getenv("OPENAI_API_KEY")

def create_word_doc(title, sections, references, topic):
    doc = Document()
    doc.add_heading(title, level=1)
    
    section_order = [
        "Abstract",
        "Introduction",
        "Literature Review",
        "Background Work",
        "Proposed Methodology",
        "Results",
        "Conclusion",
        "References"
    ]
    
    for section_title in section_order:
        if section_title in sections and sections[section_title].strip():
            doc.add_heading(section_title, level=2)
            doc.add_paragraph(sections[section_title].strip())
    
    if references.strip():
        doc.add_heading("References", level=2)
        for idx, ref in enumerate(references.split('\n')):
            if ref.strip():
                doc.add_paragraph(f"[{idx+1}] {ref.strip()}", style='List Number')
    
    doc.save(f"Research_Article_on_{topic.replace(' ', '_')}.docx")

supervisor = Agent(
    role="Supervisor",
    goal="Review and provide final approval for the research article.",
    backstory="You are the supervisor responsible for ensuring the research article meets the highest standards.",
    allow_delegation=False,
    verbose=True
)

co_supervisor = Agent(
    role="Co-Supervisor",
    goal="Guide the research process and ensure the quality of the research conducted by the students.",
    backstory="You oversee the students' work, provide detailed feedback, and ensure the research progresses as planned.",
    allow_delegation=True,
    verbose=True
)

student = Agent(
    role="Student",
    goal="Conduct research and write specific sections of the article based on the provided topic.",
    backstory="You are a student focusing on the detailed research and writing of the given topic.",
    allow_delegation=False,
    verbose=True
)

def create_tasks(topic, research_gap, proposed_methodology):
    return {
        "Title": Task(
            description=f"Generate a detailed and specific title for a research article on {topic} addressing the research gap '{research_gap}' with the proposed methodology '{proposed_methodology}'.",
            expected_output=f"A detailed and specific title for the research article on {topic}, addressing the research gap '{research_gap}' with the proposed methodology '{proposed_methodology}'.",
            agent=student
        ),
        "Literature Review": Task(
            description=f"Conduct a detailed literature review on the topic of {topic}. Include a maximum of 10 related articles with proper citations.",
            expected_output=f"Detailed literature review with 10 related articles on {topic}, including proper citations and identification of research gaps.",
            agent=student
        ),
        "Proposed Methodology": Task(
            description=f"Develop a unique proposed methodology to address the research gap: {research_gap}. Describe in detail the methods, algorithms, and processes to address the selected research gap. Include diagrams, figures, equations, and tables where necessary.",
            expected_output=f"Proposed methodology to address the research gap: {research_gap}, including proper citations and visual representations.",
            agent=student
        ),
        "Results": Task(
            description=f"Write the results section comparing the proposed results with related work on {topic}. Include detailed data, analysis, visual representations like charts and graphs, and comparative tables.",
            expected_output=f"Completed results section on {topic}, including proper citations and visual representations.",
            agent=student
        ),
        "Introduction": Task(
            description=f"Write the introduction part of the paper on {topic}. Provide background information, state the research problem, and outline the significance of the study using proper references.",
            expected_output=f"Completed introduction section on {topic}, including proper citations.",
            agent=student
        ),
        "Background Work": Task(
            description=f"Write the background work section of the paper on {topic}. Include detailed background information and foundational concepts.",
            expected_output=f"Completed background work section on {topic}, including proper citations.",
            agent=student
        ),
        "Conclusion": Task(
            description=f"Write the conclusion part of the paper on {topic}. Summarize the findings, implications, and future directions of the research.",
            expected_output=f"Completed conclusion section on {topic}, including proper citations.",
            agent=student
        ),
        "Abstract": Task(
            description=f"Write a detailed abstract of the paper on {topic}. Summarize the purpose, methodology, main findings, and significance of the research.",
            expected_output=f"Completed abstract section on {topic}, including proper citations.",
            agent=student
        ),
        "References": Task(
            description=f"List all the references cited in the research article on {topic}. Format the references according to academic standards.",
            expected_output=f"Completed references section.",
            agent=student
        ),
    }

def process_section(topic, section, task):
    print(f"Processing section: {section}")
    
    crew = Crew(
        agents=[student, co_supervisor, supervisor],
        tasks=[task],
        verbose=2
    )
    
    result = crew.kickoff(inputs={"topic": topic})
    print(f"{section} Content:\n{result.strip()}\n")
    
    feedback_crew = Crew(
        agents=[co_supervisor, supervisor],
        tasks=[
            Task(
                description=f"Review and provide feedback for the {section} section.",
                expected_output=f"Reviewed and approved {section} section with feedback.",
                agent=co_supervisor
            ),
            Task(
                description=f"Final review and approval of the {section} section.",
                expected_output=f"Final approved {section} section.",
                agent=supervisor
            )
        ],
        verbose=2
    )
    
    feedback_result = feedback_crew.kickoff(inputs={"section_content": result.strip()})
    print(f"{section} Feedback:\n{feedback_result.strip()}\n")
    
    return feedback_result.strip()

def main():
    topic = input("Enter the research topic: ")
    
    # Process the literature review to identify research gaps
    literature_review_task = Task(
        description=f"Conduct a detailed literature review on the topic of {topic}. Include a maximum of 10 related articles with proper citations. Identify and summarize key research gaps.",
        expected_output=f"Detailed literature review with 10 related articles on {topic}, including proper citations and identification of research gaps.",
        agent=student
    )
    
    research_gaps = process_section(topic, "Literature Review", literature_review_task)
    research_gap_list = [gap for gap in research_gaps.split('\n') if gap.strip()]
    selected_research_gap = research_gap_list[0] if research_gap_list else "No research gap identified"
    
    print(f"Selected Research Gap: {selected_research_gap}")

    proposed_methodology_task = Task(
        description=f"Develop a unique proposed methodology to address the research gap: {selected_research_gap}. Describe in detail the methods, algorithms, and processes to address the selected research gap. Include diagrams, figures, equations, and tables where necessary.",
        expected_output=f"Proposed methodology to address the research gap: {selected_research_gap}, including proper citations and visual representations.",
        agent=student
    )
    proposed_methodology = process_section(topic, "Proposed Methodology", proposed_methodology_task)
    
    tasks = create_tasks(topic, selected_research_gap, proposed_methodology)
    
    sections = {
        "Title": "",
        "Abstract": "",
        "Introduction": "",
        "Literature Review": research_gaps,  # Use the processed literature review
        "Background Work": "",
        "Proposed Methodology": proposed_methodology,  # Use the processed proposed methodology
        "Results": "",
        "Conclusion": "",
        "References": ""
    }
    
    references = ""
    
    for section, task in tasks.items():
        if section not in ["Literature Review", "Proposed Methodology"]:  # Skip already processed sections
            sections[section] = process_section(topic, section, task)
        
        if section == "Literature Review" or section == "References":
            references += sections[section] + "\n"
        
        print(f"{section} Content:\n{sections[section]}\n")

    create_word_doc(sections["Title"], sections, references, topic)
    print("Research article has been generated and saved as final_article.docx")

if __name__ == "__main__":
    main()

